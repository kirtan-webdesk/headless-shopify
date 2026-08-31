#!/usr/bin/env python3
"""
instruction-log.py — v1.5.3
One-command instruction log generator. Reads Claude Code JSONL transcripts and
emits a chronological list of every user instruction.

Replaces the manual multi-step JSONL parsing that Kitchen Blockers Session 2
required (Gap 8). Per Skill Verification Report A3.

Usage:
    python3 instruction-log.py --project-dir ./projects/[client]
    python3 instruction-log.py --project-dir ./projects/[client] --format md > log.md
    python3 instruction-log.py --project-dir ./projects/[client] --format docx --out log.docx
    python3 instruction-log.py --session-file path/to/session.jsonl

Output formats:
    text  (default) — readable text, one instruction per line with timestamp
    md           — markdown table
    docx         — Word document table (requires python-docx)
    json         — structured JSON for further processing
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime


def extract_user_messages(jsonl_path: Path) -> list:
    """Extract all user-role messages from a JSONL transcript."""
    instructions = []

    if not jsonl_path.exists():
        return instructions

    line_num = 0
    with open(jsonl_path) as f:
        for line in f:
            line_num += 1
            line = line.strip()
            if not line:
                continue
            try:
                entry = json.loads(line)
            except json.JSONDecodeError:
                continue

            # Different JSONL formats place the user message in different fields
            text = None
            timestamp = entry.get("timestamp") or entry.get("created_at") or ""

            # Format 1: {"role": "user", "content": "..."}
            if entry.get("role") == "user":
                content = entry.get("content")
                if isinstance(content, str):
                    text = content
                elif isinstance(content, list):
                    parts = [
                        c.get("text", "") for c in content
                        if isinstance(c, dict) and c.get("type") == "text"
                    ]
                    text = "\n".join(p for p in parts if p)

            # Format 2: {"type": "user_message", "message": {...}}
            elif entry.get("type") in ("user_message", "human"):
                msg = entry.get("message")
                if isinstance(msg, dict):
                    text = msg.get("text") or msg.get("content")
                elif isinstance(msg, str):
                    text = msg

            # Format 3: {"user": {"text": "..."}}
            elif "user" in entry and isinstance(entry["user"], dict):
                text = entry["user"].get("text")

            if text and isinstance(text, str):
                # Skip system reminders / tool outputs masquerading as user
                if text.startswith("<system-reminder>") or text.startswith("<tool_result>"):
                    continue
                instructions.append({
                    "line": line_num,
                    "timestamp": timestamp,
                    "text": text.strip(),
                    "char_count": len(text),
                })

    return instructions


def find_jsonl_files(project_dir: Path) -> list:
    """Find all JSONL session transcripts in the project."""
    candidates = []
    for pattern in ["sessions/*.jsonl", "transcripts/*.jsonl", "*.jsonl"]:
        candidates.extend(project_dir.glob(pattern))
    return sorted(set(candidates))


def format_text(instructions: list) -> str:
    lines = []
    for i, inst in enumerate(instructions, 1):
        ts = inst.get("timestamp", "")
        snippet = inst["text"].replace("\n", " ")
        if len(snippet) > 200:
            snippet = snippet[:197] + "..."
        lines.append(f"{i:4d}. [{ts}] {snippet}")
    return "\n".join(lines)


def format_markdown(instructions: list) -> str:
    lines = ["# Instruction Log", "", f"Generated: {datetime.utcnow().isoformat()}Z", "",
             f"Total instructions: {len(instructions)}", "",
             "| # | Timestamp | Instruction |", "|---|-----------|-------------|"]
    for i, inst in enumerate(instructions, 1):
        ts = inst.get("timestamp", "")
        snippet = inst["text"].replace("\n", " ").replace("|", "\\|")
        if len(snippet) > 250:
            snippet = snippet[:247] + "..."
        lines.append(f"| {i} | {ts} | {snippet} |")
    return "\n".join(lines)


def format_json(instructions: list) -> str:
    return json.dumps({
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "total": len(instructions),
        "instructions": instructions,
    }, indent=2)


def format_docx(instructions: list, out_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    doc.add_heading("Instruction Log", level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
    doc.add_paragraph(f"Total instructions: {len(instructions)}")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "#"
    header[1].text = "Timestamp"
    header[2].text = "Instruction"

    for i, inst in enumerate(instructions, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = inst.get("timestamp", "")
        text = inst["text"]
        if len(text) > 500:
            text = text[:497] + "..."
        row[2].text = text

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Generate instruction log from Claude Code JSONL transcripts.")
    parser.add_argument("--project-dir", help="Path to project workspace")
    parser.add_argument("--session-file", help="Specific JSONL file to parse (overrides project-dir)")
    parser.add_argument("--format", default="text", choices=["text", "md", "json", "docx"])
    parser.add_argument("--out", help="Output file path (required for docx, optional for others)")
    args = parser.parse_args()

    # Collect JSONL files
    if args.session_file:
        files = [Path(args.session_file)]
    elif args.project_dir:
        files = find_jsonl_files(Path(args.project_dir))
        if not files:
            print(f"ERROR: no JSONL files found in {args.project_dir}", file=sys.stderr)
            sys.exit(1)
    else:
        parser.error("--project-dir or --session-file required")

    # Extract instructions from all files, sorted by timestamp (or line if no ts)
    all_instructions = []
    for f in files:
        all_instructions.extend(extract_user_messages(f))

    # Sort by timestamp if available, else preserve insertion order
    all_instructions.sort(key=lambda x: (x.get("timestamp") or "", x.get("line", 0)))

    # Format
    if args.format == "text":
        output = format_text(all_instructions)
    elif args.format == "md":
        output = format_markdown(all_instructions)
    elif args.format == "json":
        output = format_json(all_instructions)
    elif args.format == "docx":
        if not args.out:
            print("ERROR: --out required for docx format", file=sys.stderr)
            sys.exit(1)
        format_docx(all_instructions, Path(args.out))
        print(f"Wrote {len(all_instructions)} instructions to {args.out}", file=sys.stderr)
        return

    # Print or write
    if args.out:
        Path(args.out).write_text(output)
        print(f"Wrote {len(all_instructions)} instructions to {args.out}", file=sys.stderr)
    else:
        print(output)


if __name__ == "__main__":
    main()
